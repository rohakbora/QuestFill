import io
from datetime import datetime


def build_export(fmt: str, name: str, pairs: list[dict]) -> tuple[bytes, str]:
    fmt = fmt.lower()
    if fmt == "docx": return _docx(name, pairs), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "xlsx": return _xlsx(name, pairs), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if fmt == "pdf":  return _pdf(name, pairs),  "application/pdf"
    raise ValueError(f"Unsupported format: {fmt}")


def _docx(name: str, pairs: list[dict]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph()

    for item in sorted(pairs, key=lambda x: x["order_index"]):
        p = doc.add_paragraph()
        r = p.add_run(f"Q{item['order_index'] + 1}. {item['question']}")
        r.bold = True

        doc.add_paragraph(item["answer"])

        if item.get("citation"):
            cp = doc.add_paragraph()
            cr = cp.add_run(f"Source: {item['citation']}")
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx(name: str, pairs: list[dict]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Questionnaire"
    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 60
    ws.column_dimensions["C"].width = 70
    ws.column_dimensions["D"].width = 35
    ws.column_dimensions["E"].width = 12

    hfill = PatternFill("solid", fgColor="1C1C1F")
    hfont = Font(name="Calibri", bold=True, color="F5F5F5", size=11)

    for col, header in enumerate(["#", "Question", "Answer", "Citation", "Confidence"], 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = hfill    # ← was: cell.font = hfill; cell.fill = hfill; cell.font = hfont
        cell.font = hfont
        cell.alignment = Alignment(horizontal="center")

    for i, item in enumerate(sorted(pairs, key=lambda x: x["order_index"]), start=2):
        conf = item.get("confidence", 0)
        for col, val in enumerate([
            item["order_index"] + 1, item["question"], item["answer"],
            item.get("citation") or "—", f"{conf:.0%}" if conf > 0 else "N/A",
        ], 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.font = Font(name="Calibri", size=10)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _pdf(name: str, pairs: list[dict]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    base = getSampleStyleSheet()

    title_s = ParagraphStyle("T", parent=base["Title"],  fontSize=18, spaceAfter=4)
    q_s     = ParagraphStyle("Q", parent=base["Normal"], fontSize=11, fontName="Helvetica-Bold", spaceAfter=4)
    a_s     = ParagraphStyle("A", parent=base["Normal"], fontSize=10, leftIndent=12, spaceAfter=3)
    m_s     = ParagraphStyle("M", parent=base["Normal"], fontSize=8,  leftIndent=12, textColor=colors.grey)

    story = [
        Paragraph(name, title_s),
        Paragraph(f"Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", m_s),
        Spacer(1, 0.4*cm),
    ]

    for item in sorted(pairs, key=lambda x: x["order_index"]):
        story += [
            Paragraph(f"Q{item['order_index'] + 1}. {item['question']}", q_s),
            Paragraph(item["answer"], a_s),
        ]
        if item.get("citation"):
            story.append(Paragraph(f"Source: {item['citation']}", m_s))
        story += [Spacer(1, 0.3*cm), HRFlowable(width="100%", color=colors.lightgrey), Spacer(1, 0.3*cm)]

    doc.build(story)
    return buf.getvalue()
